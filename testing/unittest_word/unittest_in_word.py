import unittest
from docx import Document
from docx.shared import Pt


# Zu testende Funktion
def add(a, b):
    return a + b


# Tests
class TestAdd(unittest.TestCase):

    def test_add1(self):
        self.assertEqual(add(2, 3), 5)

    def test_add2(self):
        self.assertEqual(add(10, 5), 15)

    def test_add3(self):
        self.assertEqual(add(3, 3), 7)   # Fehler


# Eigene TestResult-Klasse
class WordTestResult(unittest.TextTestResult):
    pass


#############################################################
suite = unittest.defaultTestLoader.loadTestsFromTestCase(TestAdd)

runner = unittest.TextTestRunner(
    resultclass=WordTestResult,
    verbosity=2
)
result = runner.run(suite)

#############################################################
print(result)

# Word-Dokument erstellen
doc = Document()

titel = doc.add_heading("Unittest-Ergebnisse", level=1)
titel.runs[0].font.size = Pt(18)


doc.add_paragraph(f"Ausgeführte Tests: {result.testsRun}")
doc.add_paragraph(f"Fehler: {len(result.errors)}")
doc.add_paragraph(f"Fehlgeschlagen: {len(result.failures)}")
doc.add_paragraph(f"Erfolgreich: {result.testsRun - len(result.errors) - len(result.failures)}")

##################
doc.add_heading("Fehlgeschlagene Tests", level=2)
if result.failures:
    for test, traceback in result.failures:
        doc.add_heading(str(test), level=3)
        doc.add_paragraph(traceback)
else:
    doc.add_paragraph("Keine.")

##################
doc.add_heading("Fehler", level=2)
if result.errors:
    for test, traceback in result.errors:
        doc.add_heading(str(test), level=3)
        doc.add_paragraph(traceback)
else:
    doc.add_paragraph("Keine.")







doc.save(r"D:\OneDrive\Kurse\COMCAVE\COMCAVE - 2026KW28 - FSM 2\Teilnehmer\python_cc_fsm2_kw28\testing\unittest_word\testergebnisse.docx")

